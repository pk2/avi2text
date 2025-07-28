# -*- coding: utf-8 -*-

import os
import sys
import tempfile
import argparse
import moviepy.editor as mp
import torch
import whisperx
from whisperx.diarize import DiarizationPipeline
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_UNDERLINE
import json
import pickle
from dotenv import load_dotenv
import logging
import language_tool_python
import difflib

def zapisz_z_sledzeniem_zmian(paragraph, original_text, corrected_text):
    """
    Porównuje dwa teksty i zapisuje je do akapitu, symulując śledzenie zmian.
    Tekst usunięty jest przekreślony, a dodany podkreślony.
    """
    matcher = difflib.SequenceMatcher(None, original_text, corrected_text)
    
    for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
        if opcode == 'equal':
            paragraph.add_run(original_text[i1:i2])
        elif opcode == 'delete':
            run = paragraph.add_run(original_text[i1:i2])
            run.font.strike = True
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif opcode == 'insert':
            run = paragraph.add_run(corrected_text[j1:j2])
            run.font.underline = WD_UNDERLINE.SINGLE
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif opcode == 'replace':
            run_del = paragraph.add_run(original_text[i1:i2])
            run_del.font.strike = True
            run_del.font.color.rgb = RGBColor(255, 0, 0)
            
            run_ins = paragraph.add_run(corrected_text[j1:j2])
            run_ins.font.underline = WD_UNDERLINE.SINGLE
            run_ins.font.color.rgb = RGBColor(0, 128, 0)

def transkrybuj_i_rozpoznaj_mowcow(
    sciezka_pliku_wideo: str,
    sciezka_pliku_docx: str,
    liczba_mowcow: int,
    model_whisper: str,
    jezyk: str,
    batch_size: int,
    compute_type: str,
    asr_options: dict
):
    """
    Wyodrębnia dźwięk, transkrybuje, dzieli na mówców, przeprowadza korektę
    i zapisuje wynik do .docx z wizualnym śledzeniem zmian.
    """
    # --- Konfiguracja logowania ---
    logging.getLogger('pytorch_lightning').setLevel(logging.ERROR)
    logging.getLogger('pyannote').setLevel(logging.ERROR)
    logging.getLogger('torch').setLevel(logging.ERROR)
    logging.getLogger('language_tool_python').setLevel(logging.ERROR)

    print(f"--- Rozpoczynanie transkrypcji pliku: {sciezka_pliku_wideo} ---")

    # --- Krok 1: Konfiguracja folderu roboczego ---
    nazwa_pliku_bazowa = os.path.splitext(os.path.basename(sciezka_pliku_wideo))[0]
    folder_roboczy = f"{nazwa_pliku_bazowa}_work"
    os.makedirs(folder_roboczy, exist_ok=True)
    print(f"Używam folderu roboczego: {folder_roboczy}")

    sciezka_pliku_audio = os.path.join(folder_roboczy, "audio.flac")
    sciezka_transkrypcji = os.path.join(folder_roboczy, "transkrypcja.json")
    sciezka_aligned = os.path.join(folder_roboczy, "aligned.json")
    sciezka_diarization = os.path.join(folder_roboczy, "diarization.pkl")
    sciezka_wyniku_finalnego = os.path.join(folder_roboczy, "wynik_finalny.json")

    # --- Krok 2: Wyodrębnienie audio ---
    if not os.path.exists(sciezka_pliku_audio):
        print("Krok 1/8: Wyodrębnianie ścieżki audio...")
        if not os.path.exists(sciezka_pliku_wideo):
            print(f"BŁĄD: Plik '{sciezka_pliku_wideo}' nie został znaleziony.")
            sys.exit(1)
        try:
            wideo = mp.VideoFileClip(sciezka_pliku_wideo)
            if wideo.audio is None:
                print(f"BŁĄD: Plik wideo '{sciezka_pliku_wideo}' nie zawiera ścieżki audio.")
                sys.exit(1)
            wideo.audio.write_audiofile(sciezka_pliku_audio, codec='flac', logger=None)
            print(f"Audio zostało pomyślnie zapisane w formacie FLAC: {sciezka_pliku_audio}")
        except Exception as e:
            print(f"BŁĄD podczas przetwarzania wideo: {e}")
            sys.exit(1)
    else:
        print("Krok 1/8: Pomijanie ekstrakcji audio (plik już istnieje).")

    # --- Krok 3: Konfiguracja urządzenia ---
    device = "cuda" if torch.cuda.is_available() else "cpu"
    print(f"Krok 2/8: Używane urządzenie: {device}")
    if device == "cpu":
        print("OSTRZEŻENIE: Brak GPU. Przetwarzanie na CPU będzie znacznie wolniejsze.")
    
    # --- Krok 4: Transkrypcja ---
    if not os.path.exists(sciezka_transkrypcji):
        print(f"Krok 3/8: Ładowanie modelu WhisperX ('{model_whisper}')...")
        model = None
        try:
            model = whisperx.load_model(model_whisper, device, compute_type=compute_type, asr_options=asr_options)
            audio = whisperx.load_audio(sciezka_pliku_audio)
            print("Krok 4/8: Rozpoczynanie transkrypcji (to może potrwać)...")
            wynik_transkrypcji = model.transcribe(audio, batch_size=batch_size, language=jezyk, print_progress=True)
            with open(sciezka_transkrypcji, 'w', encoding='utf-8') as f:
                json.dump(wynik_transkrypcji, f, ensure_ascii=False, indent=4)
            print(f"Transkrypcja zakończona i zapisana w: {sciezka_transkrypcji}")
        except Exception as e:
            print(f"BŁĄD podczas transkrypcji: {e}")
            sys.exit(1)
        finally:
            if model is not None: del model
    else:
        print("Kroki 3/8 i 4/8: Pomijanie transkrypcji (pliki już istnieją).")
        with open(sciezka_transkrypcji, 'r', encoding='utf-8') as f:
            wynik_transkrypcji = json.load(f)

    # --- Krok 5: Podział na mówców ---
    if not os.path.exists(sciezka_wyniku_finalnego):
        model_a, diarize_model = None, None
        try:
            if not os.path.exists(sciezka_aligned):
                print("Krok 5/8: Wyrównywanie transkrypcji...")
                model_a, metadata = whisperx.load_align_model(language_code=wynik_transkrypcji["language"], device=device)
                wynik_aligned = whisperx.align(wynik_transkrypcji["segments"], model_a, metadata, whisperx.load_audio(sciezka_pliku_audio), device, return_char_alignments=False)
                with open(sciezka_aligned, 'w', encoding='utf-8') as f: json.dump(wynik_aligned, f, ensure_ascii=False, indent=4)
                print(f"Wyrównywanie zakończone i zapisane w: {sciezka_aligned}")
                del model_a
            else:
                print("Krok 5/8: Pomijanie wyrównywania (plik już istnieje).")
                with open(sciezka_aligned, 'r', encoding='utf-8') as f: wynik_aligned = json.load(f)

            if not os.path.exists(sciezka_diarization):
                print("Krok 6/8: Rozpoznawanie mówców (diarization)...")
                hf_token = os.environ.get("HUGGING_FACE_TOKEN")
                if hf_token is None:
                    print("\nBŁĄD KRYTYCZNY: Brak tokena HUGGING_FACE_TOKEN.")
                    sys.exit(1)
                diarize_model = DiarizationPipeline(use_auth_token=hf_token, device=device)
                diarize_segments = diarize_model(sciezka_pliku_audio, min_speakers=liczba_mowcow, max_speakers=liczba_mowcow)
                with open(sciezka_diarization, 'wb') as f: pickle.dump(diarize_segments, f)
                print(f"Diarization zakończony i zapisany w: {sciezka_diarization}")
                del diarize_model
            else:
                print("Krok 6/8: Pomijanie diarization (plik już istnieje).")
                with open(sciezka_diarization, 'rb') as f: diarize_segments = pickle.load(f)

            wynik_finalny = whisperx.assign_word_speakers(diarize_segments, wynik_aligned)
            with open(sciezka_wyniku_finalnego, 'w', encoding='utf-8') as f: json.dump(wynik_finalny, f, ensure_ascii=False, indent=4)
            print(f"Przypisano mówców i zapisano wynik w: {sciezka_wyniku_finalnego}")
        except Exception as e:
            print(f"\nBŁĄD podczas rozpoznawania mówców: {e}")
            sys.exit(1)
    else:
        print("Kroki 5/8 i 6/8: Pomijanie (finalny plik z mówcami już istnieje).")
        with open(sciezka_wyniku_finalnego, 'r', encoding='utf-8') as f:
            wynik_finalny = json.load(f)

    # --- Krok 7: Zaawansowana korekta tekstu ---
    print("Krok 7/8: Inicjowanie zaawansowanej korekty gramatycznej (LanguageTool)...")
    tool = None
    try:
        lang_map_lt = {'pl': 'pl-PL', 'en': 'en-US', 'de': 'de-DE'}
        lt_lang_code = lang_map_lt.get(jezyk, jezyk)
        tool = language_tool_python.LanguageTool(lt_lang_code)

    except Exception as e:
        print(f"BŁĄD: Nie udało się zainicjować LanguageTool. Upewnij się, że masz zainstalowaną Javę. Błąd: {e}")

    # --- Krok 8: Zapis do pliku Word ---
    print(f"Krok 8/8: Zapisywanie wyniku do pliku {sciezka_pliku_docx}...")
    try:
        doc = Document()
        doc.add_heading(f'Transkrypcja pliku: {os.path.basename(sciezka_pliku_wideo)}', 0)
        
        aktualny_mowca = None
        aktualna_kwestia = []

        for segment in wynik_finalny["segments"]:
            if "speaker" not in segment: segment['speaker'] = "MÓWCA_NIEZNANY"
            segment_speaker = segment["speaker"].replace("SPEAKER_", "Mówca ")

            if aktualny_mowca != segment_speaker and aktualny_mowca is not None:
                tekst_do_korekty = ''.join(aktualna_kwestia).lstrip()
                
                if tool:
                    poprawiony_tekst = tool.correct(tekst_do_korekty)
                else:
                    poprawiony_tekst = tekst_do_korekty
                
                p = doc.add_paragraph()
                p.add_run(f"{aktualny_mowca}: ").bold = True
                zapisz_z_sledzeniem_zmian(p, tekst_do_korekty, poprawiony_tekst)
                aktualna_kwestia = []
            
            aktualny_mowca = segment_speaker
            
            if 'words' in segment:
                for word_info in segment['words']:
                    aktualna_kwestia.append(word_info['word'] + " ")
            else:
                aktualna_kwestia.append(segment.get('text', '').strip() + " ")

        if aktualny_mowca is not None and aktualna_kwestia:
            tekst_do_korekty = ''.join(aktualna_kwestia).lstrip()
            if tool:
                poprawiony_tekst = tool.correct(tekst_do_korekty)
            else:
                poprawiony_tekst = tekst_do_korekty
            
            p = doc.add_paragraph()
            p.add_run(f"{aktualny_mowca}: ").bold = True
            zapisz_z_sledzeniem_zmian(p, tekst_do_korekty, poprawiony_tekst)

        doc.save(sciezka_pliku_docx)
        print("\n--- Zakończono pomyślnie! ---")
        print(f"Wynik został zapisany w pliku: {sciezka_pliku_docx}")

    except Exception as e:
        print(f"BŁĄD podczas zapisu do pliku .docx: {e}")
    finally:
        if tool:
            tool.close()

if __name__ == "__main__":
    load_dotenv()
    default_compute_type = "float16" if torch.cuda.is_available() else "int8"

    parser = argparse.ArgumentParser(
        description="Zaawansowana transkrypcja wideo z podziałem na mówców i zapisem do Word.",
        formatter_class=argparse.RawTextHelpFormatter
    )
    parser.add_argument("sciezka_wideo", type=str, help="Ścieżka do pliku wideo.")
    parser.add_argument("--liczba_mowcow", type=int, default=os.getenv("DEFAULT_SPEAKERS", None), help="Liczba mówców.")
    parser.add_argument("--plik_wyjsciowy", type=str, default=None, help="Nazwa pliku .docx.")
    parser.add_argument("--model", type=str, default=os.getenv("DEFAULT_MODEL", "large-v2"), choices=["tiny", "base", "small", "medium", "large-v1", "large-v2", "large-v3"], help="Model Whisper do użycia.")
    parser.add_argument("--jezyk", type=str, default=os.getenv("DEFAULT_LANGUAGE", "pl"), help="Kod języka (np. 'pl', 'en').")
    parser.add_argument("--batch_size", type=int, default=16, help="Liczba segmentów przetwarzanych równolegle.")
    parser.add_argument("--cpu_threads", type=int, default=os.cpu_count(), help="Liczba wątków CPU do użycia.")
    parser.add_argument("--compute_type", type=str, default=default_compute_type, choices=["float16", "float32", "int8", "int8_float16"], help=f"Typ obliczeń. Domyślnie: '{default_compute_type}'.")
    parser.add_argument("--beam_size", type=int, default=5, help="Liczba 'promieni' w beam search.")
    
    args = parser.parse_args()

    if args.liczba_mowcow is None:
        print("BŁĄD: Liczba mówców jest wymagana.")
        sys.exit(1)
        
    plik_wyjsciowy = args.plik_wyjsciowy
    if plik_wyjsciowy is None:
        nazwa_bazowa = os.path.splitext(os.path.basename(args.sciezka_wideo))[0]
        plik_wyjsciowy = f"{nazwa_bazowa}.docx"
        
    if not torch.cuda.is_available():
        print(f"Ustawiam liczbę wątków CPU na: {args.cpu_threads}")
        torch.set_num_threads(args.cpu_threads)

    asr_options = {"beam_size": args.beam_size}

    transkrybuj_i_rozpoznaj_mowcow(
        args.sciezka_wideo,
        plik_wyjsciowy,
        int(args.liczba_mowcow),
        args.model,
        args.jezyk,
        args.batch_size,
        args.compute_type,
        asr_options
    )

